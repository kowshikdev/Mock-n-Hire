"""Resume text extraction, shared by the recruiter and candidate paths.

Both sides used to do their own thing: the recruiter path used pdfminer plus
python-docx, the candidate path used PyMuPDF and rejected anything that was
not a PDF. Neither recovered reading order on multi-column layouts, and
neither noticed when a PDF yielded no text at all.
"""

from __future__ import annotations

import io
import logging

logger = logging.getLogger(__name__)

# Below this, a "successful" extraction is almost certainly a scanned or
# image-only PDF. Interviewing someone on an empty string produces generic
# questions with no connection to them at all, so callers are expected to
# treat this as a hard failure rather than carry on.
MIN_USABLE_CHARS = 100

SUPPORTED_EXTENSIONS = (".pdf", ".docx")

# A gutter narrower than this is just word spacing, not a column boundary.
MIN_GUTTER_WIDTH = 18.0

# Only look for a gutter in the middle of the page. A resume with a narrow
# left rail still splits well within this range, and it stops a ragged right
# margin from being mistaken for one.
GUTTER_SEARCH_RANGE = (0.22, 0.78)


class UnreadableResume(Exception):
    """Raised when a document yields no usable text."""


def _page_lines(page) -> list[tuple[float, float, float, str]]:
    """Every text line on a page as (x0, y0, x1, text).

    Lines, not blocks. PyMuPDF's block segmentation merges text that shares a
    baseline, so on a two-column resume a single block comes back holding
    *both* columns' text for that row -- 'SKILLS\\nEXPERIENCE'. No amount of
    sorting at block level can separate those again; the bboxes needed to tell
    the columns apart only exist one level down.
    """
    lines = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:  # 0 = text, 1 = image
            continue
        for line in block.get("lines", []):
            text = "".join(span["text"] for span in line.get("spans", [])).strip()
            if not text:
                continue
            x0, y0, x1, _ = line["bbox"]
            lines.append((x0, y0, x1, text))
    return lines


def _find_gutter(lines: list[tuple[float, float, float, str]], page_width: float) -> float | None:
    """Locate the vertical whitespace channel between two columns, if any.

    Scans candidate x positions across the middle of the page and counts how
    many lines straddle each one, then takes the widest run at the minimum
    crossing count.

    That minimum is deliberately not required to be zero. Almost every
    two-column resume has at least one line that legitimately spans the whole
    width -- the name header, a section banner, a horizontal rule -- and
    demanding a completely clear channel let a single wide header suppress
    column detection for the entire page.
    """
    if len(lines) < 8:
        return None

    lo = int(page_width * GUTTER_SEARCH_RANGE[0])
    hi = int(page_width * GUTTER_SEARCH_RANGE[1])
    if hi <= lo:
        return None

    counts = [
        sum(1 for x0, _, x1, _ in lines if x0 < x < x1)
        for x in range(lo, hi + 1)
    ]

    # Positions crossed by no more than a handful of lines are gutter
    # candidates. The allowance is what lets the name header and section
    # banners span the page without hiding the channel underneath them.
    tolerance = max(1, int(len(lines) * 0.15))

    runs: list[tuple[int, int]] = []
    run_start: int | None = None
    for i, count in enumerate(counts):
        if count <= tolerance:
            if run_start is None:
                run_start = i
        elif run_start is not None:
            runs.append((run_start, i))
            run_start = None
    if run_start is not None:
        runs.append((run_start, len(counts)))

    # Score every candidate rather than taking the widest. The widest run on
    # a two-column page is usually the empty right *margin* beyond the longest
    # line, which splits the page into "everything" and "nothing" -- so
    # balance decides, and width only breaks ties.
    best: tuple[int, int, float] | None = None
    for start, end in runs:
        if (end - start) < MIN_GUTTER_WIDTH:
            continue
        split = lo + (start + end) / 2
        left = sum(1 for _, _, x1, _ in lines if x1 <= split)
        right = sum(1 for x0, _, _, _ in lines if x0 >= split)
        if left < 3 or right < 3:
            continue
        score = (min(left, right), end - start, split)
        if best is None or score > best:
            best = score

    return best[2] if best else None


def _order_lines(lines: list[tuple[float, float, float, str]], page_width: float) -> list[str]:
    """Put a page's lines into human reading order."""
    if not lines:
        return []

    split = _find_gutter(lines, page_width)
    if split is None:
        return [text for _, _, _, text in sorted(lines, key=lambda l: (round(l[1], 1), l[0]))]

    # Full-width lines -- the name header, a horizontal rule, a section that
    # spans both columns -- break the page into horizontal bands. Within each
    # band the left column is read top to bottom, then the right. Without this
    # banding a mid-page full-width heading would be torn away from the
    # column content it introduces.
    ordered: list[str] = []
    pending: list[tuple[float, float, float, str]] = []

    def flush() -> None:
        if not pending:
            return
        left = sorted((l for l in pending if l[2] <= split), key=lambda l: (round(l[1], 1), l[0]))
        right = sorted((l for l in pending if l[0] >= split), key=lambda l: (round(l[1], 1), l[0]))
        ordered.extend(text for _, _, _, text in left)
        ordered.extend(text for _, _, _, text in right)
        pending.clear()

    for line in sorted(lines, key=lambda l: (round(l[1], 1), l[0])):
        x0, _, x1, text = line
        if x0 < split < x1:  # spans the gutter
            flush()
            ordered.append(text)
        else:
            pending.append(line)
    flush()

    return ordered


def _extract_pdf(content: bytes) -> str:
    import fitz  # PyMuPDF

    parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for page in doc:
            parts.extend(_order_lines(_page_lines(page), page.rect.width))
    return "\n".join(parts).strip()


def _extract_docx(content: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(content))
    parts = [p.text for p in document.paragraphs if p.text.strip()]

    # Plenty of resumes lay everything out in a borderless table, which
    # python-docx does not surface through `.paragraphs` at all -- those
    # documents extracted as an empty string.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()


def extract_resume_text(content: bytes, filename: str) -> str:
    """Extract reading-order text from a resume.

    Raises UnreadableResume for unsupported formats, corrupt files, and
    documents that yield essentially no text.
    """
    lowered = (filename or "").lower()
    if not lowered.endswith(SUPPORTED_EXTENSIONS):
        raise UnreadableResume("Only PDF and DOCX resumes are supported.")

    try:
        text = _extract_pdf(content) if lowered.endswith(".pdf") else _extract_docx(content)
    except UnreadableResume:
        raise
    except Exception as e:
        logger.warning(f"Resume extraction failed for {filename}: {e}")
        raise UnreadableResume("That file couldn't be opened. Is it a valid PDF or DOCX?")

    if len(text) < MIN_USABLE_CHARS:
        raise UnreadableResume(
            "We couldn't read any text from that file. If it's a scan or an image "
            "export, please upload a text-based PDF or a DOCX instead."
        )

    return text
